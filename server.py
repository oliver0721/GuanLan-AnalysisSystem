import os
import pickle
import re
import sys
import pandas as pd
import requests
import urllib3
from bs4 import BeautifulSoup
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 禁用安全警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# --- 引入 LiteLLM (支持所有模型) ---
from litellm import completion

# 导入你现有的功能模块
try:
    from weibo_server_zhutie import scrape_weibo_posts
    from weibo_server_process import process_weibo_data
except ImportError:
    print("❌ 错误：未找到爬虫脚本，请确保 weibo_server_zhutie.py 在同一目录下")

app = Flask(__name__)
CORS(app)

# 路径配置
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
COOKIE_FILE = os.path.join(BASE_DIR, "weibo_cookies.pkl")
INDEX_FILE = os.path.join(BASE_DIR, "index.html")


# --- 辅助函数 ---
def extract_path(result_str, pattern):
    match = re.search(pattern, result_str)
    return match.group(1).strip() if match else None


# --- 爬取微博热搜 (增强版) ---
def get_weibo_hot_search_list():
    url = "https://s.weibo.com/top/summary"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Cookie": ""
        # 匿名Cookie
    }

    try:
        # verify=False 极其重要，解决本地SSL报错
        resp = requests.get(url, headers=headers, timeout=8, verify=False)
        if resp.status_code != 200:
            return [{"title": f"连接失败 ({resp.status_code})", "link": "#"}]

        soup = BeautifulSoup(resp.text, 'html.parser')
        items = soup.select('td.td-02 a')

        hot_list = []
        for item in items:
            title = item.get_text().strip()
            href = item.get('href', '')

            # 补全链接
            if href.startswith('/'):
                link = "https://s.weibo.com" + href
            elif href.startswith('http'):
                link = href
            else:
                link = f"https://s.weibo.com/weibo?q={title}"

            # 过滤置顶广告
            if "javascript" in href or not href or not title: continue

            hot_list.append({"title": title, "link": link})

        # 修改后：从索引1开始取（跳过索引0），取到索引6（即真正的第1到第5名）
        return hot_list[1:6] if len(hot_list) > 1 else [{"title": "暂无数据", "link": "#"}]

    except Exception as e:
        print(f"热搜获取异常: {e}")
        return [{"title": "网络连通性错误 (请检查代理)", "link": "#"}]


# --- 路由 ---
@app.route('/')
def index():
    if os.path.exists(INDEX_FILE): return send_file(INDEX_FILE)
    return "错误：找不到 index.html", 404


@app.route('/files')
def get_file():
    path = request.args.get('path')
    if path and os.path.exists(path): return send_file(path)
    return "File not found", 404


@app.route('/api/hot_search', methods=['GET'])
def hot_search():
    data = get_weibo_hot_search_list()
    return jsonify({"status": "success", "data": data})


@app.route('/api/update_cookie', methods=['POST'])
def update_cookie():
    data = request.json
    try:
        cookies = []
        for item in data.get('cookie', '').split(';'):
            if '=' in item:
                k, v = item.strip().split('=', 1)
                cookies.append({'name': k, 'value': v, 'domain': '.weibo.com', 'path': '/'})
        with open(COOKIE_FILE, "wb") as f:
            pickle.dump(cookies, f)
        return jsonify({"status": "success", "count": len(cookies)})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/api/analyze', methods=['POST'])
def analyze():
    data = request.json
    keyword = data.get('keyword')
    api_key = data.get('api_key')
    model = data.get('model')  # 例如 "openai/gpt-4" 或 "deepseek/deepseek-chat"
    base_url = data.get('base_url')

    if not keyword or not api_key:
        return jsonify({"status": "error", "message": "请填写关键词和 API Key"}), 400

    try:
        # 1. 爬取
        print(f"Step 1: 爬取 {keyword}")
        scrape_res = scrape_weibo_posts(keyword, pages=10)
        csv_path = extract_path(scrape_res, r"已保存 CSV：(.*?)(?:\n|$)")
        if not csv_path or not os.path.exists(csv_path):
            return jsonify({"status": "error", "message": "爬虫未生成数据，请更新 Cookie"}), 500

        # 2. 处理
        print("Step 2: 绘图")
        proc_res = process_weibo_data(csv_path)
        img_paths = {
            'sentiment': extract_path(proc_res, r"情感分布图: (.*?)(?:\n|$)"),
            'trend': extract_path(proc_res, r"每日趋势图: (.*?)(?:\n|$)"),
            'wordcloud': extract_path(proc_res, r"词云图: (.*?)(?:\n|$)")
        }

        # 3. LLM 分析
        print(f"Step 3: LiteLLM ({model})")
        df = pd.read_csv(csv_path.replace('.csv', '_processed.csv'))
        summary = f"关键词：{keyword}，样本：{len(df)}条，正面：{len(df[df['sentiment_class'] == 'positive'])}"
        sample = df[['content', 'sentiment_class']].head(30).to_string()

        sys_prompt = "你是一名专业的舆情分析师，擅长撰写学术风格的深度报告。"
        user_prompt = f"""
        基于以下数据撰写【{keyword}】的舆情研究报告。
        数据摘要：{summary}
        评论样本：\n{sample}

        【写作要求】
        1. 风格模仿“新浪舆情通”，包含：## 舆情综述、## 情感倾向、## 传播趋势、## 研判建议。
        2. **必须在正文中插入图片占位符**（不要把图堆在最后）：
           - 情感分析章节必需包含： (此处插入情感分布图)
           - 传播趋势章节必需包含： (此处插入趋势图)
           - 舆论热词章节必需包含： (此处插入词云图)
        3. 字数 800-1000 字。
        """

        # 调用 litellm
        resp = completion(
            model=model,
            messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": user_prompt}],
            api_key=api_key,
            base_url=base_url if base_url else None
        )
        text = resp.choices[0].message.content

        # 4. 生成 Word
        print("Step 4: Word")
        doc = Document()
        doc.add_heading(f'{keyword} 舆情分析报告', 0)
        for line in text.split('\n'):
            line = line.strip()
            if not line: continue
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), 1)
            elif "此处插入情感分布图" in line and img_paths['sentiment']:
                try:
                    doc.add_picture(img_paths['sentiment'], width=Inches(5.5))
                except:
                    pass
            elif "此处插入趋势图" in line and img_paths['trend']:
                try:
                    doc.add_picture(img_paths['trend'], width=Inches(6))
                except:
                    pass
            elif "此处插入词云图" in line and img_paths['wordcloud']:
                try:
                    doc.add_picture(img_paths['wordcloud'], width=Inches(6))
                except:
                    pass
            else:
                doc.add_paragraph(line)

        doc_filename = f"{keyword}_report.docx"
        doc_path = os.path.join(os.path.dirname(csv_path), doc_filename)
        doc.save(doc_path)

        return jsonify({"status": "success", "analysis": text, "images": img_paths, "doc_path": doc_path})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": str(e)}), 500


if __name__ == '__main__':
    # 端口改为 8080 避免冲突
    app.run(host='0.0.0.0', port=8080, debug=True)